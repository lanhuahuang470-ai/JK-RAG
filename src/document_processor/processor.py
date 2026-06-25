"""
文档处理模块
支持 PDF（PyPDF2）和 Word（python-docx）文本提取，
使用固定长度切片策略（chunk_size=1000, chunk_overlap=200）
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List

import PyPDF2
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from config.settings import settings
from src.utils import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """
    企业级文档处理器

    职责：
    1. 从 PDF / DOCX 提取纯文本
    2. 清洗文本（去除噪声字符）
    3. 固定长度切片 → List[Document]
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}

    def __init__(
        self,
        chunk_size: int = settings.chunk_size,
        chunk_overlap: int = settings.chunk_overlap,
    ) -> None:
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        # 固定长度切片（回退到字符边界，保留语义连贯性）
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
        )

        logger.info(
            "DocumentProcessor 初始化完成 | "
            f"chunk_size={chunk_size} | chunk_overlap={chunk_overlap}"
        )

    # ──────────────────────────────────────────────────────────────────────────
    # 公共接口
    # ──────────────────────────────────────────────────────────────────────────

    def process_file(self, file_path: str | Path) -> List[Document]:
        """
        处理单个文件，返回切分后的 LangChain Document 列表

        Args:
            file_path: 文件路径（PDF / DOCX）

        Returns:
            List[Document]：每个 Document 包含 page_content 和 metadata
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的文件格式：{suffix}。"
                f"当前支持：{self.SUPPORTED_EXTENSIONS}"
            )

        logger.info(f"开始处理文件：{path.name}")

        if suffix == ".pdf":
            raw_text = self._extract_pdf(path)
        else:  # .docx / .doc
            raw_text = self._extract_docx(path)

        cleaned = self._clean_text(raw_text)
        chunks = self._split(cleaned, source=path.name)

        logger.info(f"文件 {path.name} 处理完成 | 共切分 {len(chunks)} 个 chunk")
        return chunks

    def process_files(self, file_paths: List[str | Path]) -> List[Document]:
        """批量处理多个文件"""
        all_docs: List[Document] = []
        for fp in file_paths:
            try:
                all_docs.extend(self.process_file(fp))
            except Exception as e:
                logger.error(f"处理文件 {fp} 失败：{e}")
        logger.info(f"批量处理完成 | 共 {len(file_paths)} 个文件 | {len(all_docs)} 个 chunk")
        return all_docs

    # ──────────────────────────────────────────────────────────────────────────
    # PDF 提取
    # ──────────────────────────────────────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> str:
        """使用 PyPDF2 逐页提取 PDF 文本"""
        texts: List[str] = []
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                total_pages = len(reader.pages)
                logger.debug(f"PDF 共 {total_pages} 页")

                for page_num, page in enumerate(reader.pages, start=1):
                    try:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            texts.append(page_text)
                    except Exception as e:
                        logger.warning(f"第 {page_num} 页提取失败：{e}")

        except PyPDF2.errors.PdfReadError as e:
            raise RuntimeError(f"PDF 读取错误（可能是加密文件）：{e}") from e

        full_text = "\n".join(texts)
        logger.debug(f"PDF 提取完成 | 原始字符数：{len(full_text)}")
        return full_text

    # ──────────────────────────────────────────────────────────────────────────
    # DOCX 提取
    # ──────────────────────────────────────────────────────────────────────────

    def _extract_docx(self, path: Path) -> str:
        """使用 python-docx 提取 Word 文档文本（含表格）"""
        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            raise RuntimeError(f"DOCX 读取错误：{e}") from e

        parts: List[str] = []

        # 提取正文段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        full_text = "\n".join(parts)
        logger.debug(f"DOCX 提取完成 | 原始字符数：{len(full_text)}")
        return full_text

    # ──────────────────────────────────────────────────────────────────────────
    # 文本清洗
    # ──────────────────────────────────────────────────────────────────────────

    @staticmethod
    def _clean_text(text: str) -> str:
        """
        文本清洗：
        - 去除控制字符（保留换行/制表符）
        - 合并多余空白行
        - 去除首尾空白
        """
        # 去除 PDF 常见噪声（页码行、水印行等简单规则）
        text = re.sub(r"\x0c", "\n", text)          # 换页符 → 换行
        text = re.sub(r"[^\S\n]+", " ", text)        # 多个空格合并
        text = re.sub(r"\n{3,}", "\n\n", text)        # 多个空行合并为双换行
        text = text.strip()
        return text

    # ──────────────────────────────────────────────────────────────────────────
    # 切分
    # ──────────────────────────────────────────────────────────────────────────

    def _split(self, text: str, source: str) -> List[Document]:
        """将清洗后的文本切分为固定长度 chunk，并注入元数据"""
        raw_chunks = self._splitter.split_text(text)
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "chunk_index": idx,
                    "chunk_size": len(chunk),
                },
            )
            for idx, chunk in enumerate(raw_chunks)
            if chunk.strip()  # 跳过空 chunk
        ]
        return documents
